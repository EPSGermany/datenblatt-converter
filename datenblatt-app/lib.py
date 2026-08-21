"""
lib.py - Kernlogik des Datenblatt-Converters (ohne GUI, ohne KI/Server)
=========================================================================
Automatische, aber immer wörtliche Übernahme von Werten aus einem fremden
PDF-Datenblatt in eine eigene Word-Vorlage. Nichts wird erfunden: jeder
eingesetzte Wert ist eine 1:1-Kopie einer Zeichenkette aus dem Quell-PDF.

Erkennungsstrategie (alles rein positions-/formatbasiert, keine KI):
- Produkttitel  = Textzeile mit der größten Schriftgröße im Dokument
- Beschreibung  = zusammenhängender Fließtext-Absatz vor der ersten
                  erkannten Wertetabelle
- Abschnitte    = kurze Textzeile (<=4 Wörter), auf die direkt eine oder
                  mehrere Wert-Paar-Zeilen folgen, wird als Abschnitts-
                  überschrift interpretiert (z.B. "Allgemeine Daten")
- Wert-Paare    = zwei durch eine auffällige horizontale Lücke getrennte
                  Wortgruppen in einer Zeile (deckt Tabellen ohne
                  sichtbare Linien/Doppelpunkt ab) ODER klassische
                  "Label: Wert"-Zeilen
- Listen        = Abschnitt, dessen "Werte" fehlen (nur einzelne kurze
                  Zeilen ohne erkennbares Paar), z.B. Lieferumfang
"""
import re
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

import pdfplumber
from docx import Document

KEYVALUE_RE = re.compile(r"^\s*([A-Za-zÄÖÜäöüß0-9\.\-\/\s]{2,40}?)\s*:\s+(.+?)\s*$")
COLUMN_GAP_THRESHOLD = 15
HEADING_MAX_WORDS = 4
NOT_FOUND = "-- nicht gefunden --"


# ---------------------------------------------------------------------------
def _parse_page_ordered(page, gap_threshold=COLUMN_GAP_THRESHOLD, page_no=0):
    """Liefert die Zeilen einer Seite in Lesereihenfolge, jede als
    ('pair', label, value, size, page_no) oder ('text', text, None, size, page_no)."""
    words = page.extract_words(extra_attrs=["size"])
    lines = defaultdict(list)
    for w in words:
        lines[round(w["top"])].append(w)

    entries = []
    for top in sorted(lines.keys()):
        row = sorted(lines[top], key=lambda w: w["x0"])
        max_size = max(w["size"] for w in row)
        if len(row) >= 2:
            gaps = [(row[i + 1]["x0"] - row[i]["x1"], i) for i in range(len(row) - 1)]
            max_gap, split_i = max(gaps, key=lambda g: g[0])
            if max_gap >= gap_threshold:
                label = " ".join(w["text"] for w in row[: split_i + 1]).strip()
                value = " ".join(w["text"] for w in row[split_i + 1 :]).strip()
                entries.append(("pair", label, value, max_size, page_no))
                continue
        m = KEYVALUE_RE.match(" ".join(w["text"] for w in row))
        if m and m.group(1).strip() and m.group(2).strip():
            entries.append(("pair", m.group(1).strip(), m.group(2).strip(), max_size, page_no))
            continue
        entries.append(("text", " ".join(w["text"] for w in row).strip(), None, max_size, page_no))
    return entries


# ---------------------------------------------------------------------------
def extract_pdf(pdf_path: str) -> dict:
    """Liest ein PDF vollständig und wörtlich aus und leitet automatisch
    Titel, Beschreibung, Abschnitte (mit Wertepaaren) und Listen ab.
    Erzeugt NIE neuen Text - nur Gruppierung von tatsächlich vorhandenem."""
    pages_entries = []  # Liste von Listen (pro Seite)
    with pdfplumber.open(pdf_path) as pdf:
        for page_no, page in enumerate(pdf.pages):
            pages_entries.append(_parse_page_ordered(page, page_no=page_no))

    # Wiederkehrende Kopf-/Fußzeilen erkennen (identischer Inhalt auf >=2
    # Seiten, z.B. "Firma XY - Datenblatt", "Serie ...", Website-Fußzeile)
    # und als Rauschen herausfiltern, damit sie die automatische
    # Abschnittserkennung nicht stören (z.B. fälschlich als neue
    # Überschrift interpretiert werden). Betrifft NICHT die inhaltliche
    # Übernahme - es werden nur Duplikate der immer gleichen Kopf-/Fußzeile
    # ignoriert, keine echten Datenwerte.
    from collections import Counter
    signature_count = Counter()
    ever_precedes_pair = set()
    for page_entries in pages_entries:
        seen_this_page = set()
        for idx, e in enumerate(page_entries):
            sig = (e[0], e[1], e[2])
            if sig not in seen_this_page:
                signature_count[sig] += 1
                seen_this_page.add(sig)
            if idx + 1 < len(page_entries):
                nxt = page_entries[idx + 1]
                # Seitenzahl-Pseudopaare ("Seite: 2"/"Page: 2") zählen NICHT
                # als "folgt ein Wertepaar" - sonst würde eine reine
                # Kopfzeile wie "EPS - Datenblatt" (der oft direkt eine
                # Seitenzahl vorausgeht) fälschlich wie eine echte
                # Abschnittsüberschrift behandelt.
                if nxt[0] == "pair" and nxt[1].strip().lower() not in ("seite", "page"):
                    ever_precedes_pair.add(sig)
    # Nur Zeilen, die NIE unmittelbar vor Wertepaaren stehen, gelten als
    # reine Kopf-/Fußzeile (Boilerplate). Eine Überschrift wie "Technical
    # data", die bei einem über den Seitenumbruch fortgesetzten
    # Datenblatt zweimal identisch auftaucht, aber jedes Mal direkt
    # Wertepaare einleitet, ist KEINE Boilerplate und bleibt erhalten.
    boilerplate = {
        sig for sig, count in signature_count.items()
        if count >= 2 and sig not in ever_precedes_pair
    }

    def _serie_prefix_len(text: str):
        """Gibt die Länge des 'Serie '/'Series '-Präfixes zurück, wenn die
        Zeile eine Serie/Series-Titelzeile ist, sonst None. Erkennt auch
        das PDF-Extraktionsartefakt fehlender Leerzeichen (z.B.
        "SeriesEPS/PUB 10000 4U")."""
        t = text.strip()
        tl = t.lower()
        if tl.startswith("serien"):  # z.B. "Seriennummer" - keine Serie-Zeile
            return None
        for prefix in ("serie", "series"):
            if tl.startswith(prefix):
                rest = t[len(prefix):]
                if rest and (rest[0] == " " or rest[0].isupper() or rest[0] == "/"):
                    return len(prefix)
        return None

    # "Serie ..." (deutsch) bzw. "Series ..." (englisch) steht meist als
    # wiederkehrende Kopfzeile auf jeder Seite (daher oben ggf. als
    # Boilerplate gefiltert) - genau daraus aber lässt sich die
    # Serienbezeichnung zuverlässig ableiten, wörtlich übernommen.
    serie = ""
    for (kind, a, b), count in signature_count.items():
        if kind != "text" or count < 2:
            continue
        plen = _serie_prefix_len(a)
        if plen is not None:
            serie = a[plen:].strip()
            break

    PAGE_NUM_RE = re.compile(r"^(seite|page)\s*[:\-]?\s*\d+$", re.IGNORECASE)

    def _is_serie_line(text: str) -> bool:
        t = text.strip()
        tl = t.lower()
        if tl.startswith("serien"):  # z.B. "Seriennummer" - KEINE Serie-Titelzeile
            return False
        for prefix in ("serie", "series"):
            if tl.startswith(prefix):
                rest = t[len(prefix):]
                if not rest:
                    return False
                # Leerzeichen danach (Normalfall) ODER Großbuchstabe/Slash
                # direkt anschließend (PDF-Extraktionsartefakt: fehlendes
                # Leerzeichen zwischen "Series" und Modellname, z.B.
                # "SeriesEPS/PUB 10000 4U")
                if rest[0] == " " or rest[0].isupper() or rest[0] == "/":
                    return True
        return False

    all_entries = []
    raw_lines = []
    for page_entries in pages_entries:
        for e in page_entries:
            sig = (e[0], e[1], e[2])
            if sig in boilerplate:
                continue
            if e[0] == "pair" and e[1].strip().lower() in ("seite", "page"):
                continue  # Seitenzahl-Artefakt, kein echter Datenwert
            if e[0] == "text" and PAGE_NUM_RE.match(e[1].strip()):
                continue  # z.B. "Page 3" als eigene Restzeile
            if e[0] == "text" and _serie_prefix_len(e[1]) is not None:
                continue  # "Serie ..."/"Series ..." - wird separat als `serie` erfasst, nie als Abschnittsüberschrift verwenden
            all_entries.append(e)
            raw_lines.append(e[1] if e[0] == "text" else f"{e[1]}: {e[2]}")

    # Titel = Textzeile mit größter Schriftgröße
    text_entries = [e for e in all_entries if e[0] == "text" and e[1]]
    titel = ""
    if text_entries:
        titel = max(text_entries, key=lambda e: e[3])[1]

    # Abschnitte: kurze Textzeile gefolgt von >=1 Paar(en) => Überschrift+Paare
    # Listen: kurze Textzeile gefolgt von kurzen Textzeilen (keine Paare) => Liste
    sections = []  # {"heading": str, "type": "pairs"/"list", "items": [...]}
    first_section_start = len(all_entries)
    i = 0
    n = len(all_entries)
    while i < n:
        kind, a, b, size, page_no = all_entries[i]
        if kind == "text" and 0 < len(a.split()) <= HEADING_MAX_WORDS:
            # schauen was folgt
            j = i + 1
            pairs, list_items = [], []
            while j < n and all_entries[j][0] == "pair":
                pairs.append([all_entries[j][1], all_entries[j][2]])
                j += 1
            if pairs:
                sections.append({"heading": a.rstrip(":"), "type": "pairs", "items": pairs})
                first_section_start = min(first_section_start, i)
                i = j
                continue
            # sonst: evtl. Liste kurzer Zeilen. Eine Liste darf über einen
            # Seitenumbruch hinweg fortgesetzt werden (manche Datenblätter
            # brechen z.B. "Lieferumfang" einfach mitten in der Liste auf
            # die nächste Seite um, ohne die Überschrift zu wiederholen) -
            # ABER nur solange die Schriftgröße zur bisherigen Liste passt.
            # Ein neuer Produkttitel auf der Folgeseite hat fast immer eine
            # deutlich andere (meist größere) Schrift und wird so nicht
            # versehentlich mit eingesammelt.
            ref_size = size
            k = i + 1
            while (
                k < n
                and all_entries[k][0] == "text"
                and 0 < len(all_entries[k][1].split()) <= HEADING_MAX_WORDS + 2
                and abs(all_entries[k][3] - ref_size) < 1.5
            ):
                if k + 1 < n and all_entries[k + 1][0] == "pair":
                    break  # diese Zeile gehört als Überschrift zu den folgenden Paaren
                list_items.append(all_entries[k][1])
                ref_size = all_entries[k][3]
                k += 1
            if list_items:
                sections.append({"heading": a.rstrip(":"), "type": "list", "items": list_items})
                first_section_start = min(first_section_start, i)
                i = k
                continue
        i += 1

    # Beschreibung = alle Textzeilen VOR dem ersten erkannten Abschnitt
    # (Lieferumfang, Allgemeine Daten, ...). Robuster als ein reiner
    # Wortanzahl-Schwellwert, da kurze Restzeilen am Satzende (z.B.
    # "Industriesteuerungen geeignet.") sonst die Beschreibung fälschlich
    # vorzeitig abschneiden würden.
    beschreibung_lines = [
        e[1] for e in all_entries[:first_section_start] if e[0] == "text"
    ]
    beschreibung = " ".join(beschreibung_lines)

    # Aufeinanderfolgende Abschnitte mit identischer Überschrift (z.B. wenn
    # eine Tabelle über einen Seitenumbruch hinweg fortgesetzt wird und die
    # Überschrift dabei wiederholt wird) zu einem Abschnitt zusammenführen.
    merged_sections = []
    for s in sections:
        if (
            merged_sections
            and merged_sections[-1]["heading"] == s["heading"]
            and merged_sections[-1]["type"] == s["type"]
        ):
            merged_sections[-1]["items"].extend(s["items"])
        else:
            merged_sections.append(s)
    sections = merged_sections

    # Produkttyp = erster Eintrag der Liste, die direkt dem Titel folgt
    # (z.B. "Labornetzgerät" direkt unter "E/PS 9040-20 T") - wörtlich
    # übernommen, keine Interpretation.
    produkttyp = ""
    for s in sections:
        if s["type"] == "list" and s["heading"] == titel and s["items"]:
            produkttyp = s["items"][0]
            break

    # Alle Wertepaare zusätzlich flach gesammelt (für Kontrolle/Suche)
    key_values = {}
    for e in all_entries:
        if e[0] == "pair":
            key_values.setdefault(e[1], e[2])

    return {
        "titel": titel,
        "serie": serie,
        "produkttyp": produkttyp,
        "beschreibung": beschreibung,
        "sections": sections,
        "key_values": key_values,
        "raw_lines": raw_lines,
    }


# ---------------------------------------------------------------------------
# Vorlage befüllen
# ---------------------------------------------------------------------------
def _replace_placeholder_text(paragraph, mapping: dict, not_found_text: str = NOT_FOUND):
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return
    new_text = full_text
    for key, val in mapping.items():
        new_text = new_text.replace("{{" + key + "}}", val if val else not_found_text)
    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def _shade_cell(cell, hex_color: str):
    """Setzt die Hintergrundfarbe einer Tabellenzelle (ohne Text zu verändern)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _fill_named_section_table(table, marker_key: str, sections: list, shade_color: str = "DCE9F7"):
    """Füllt EINE Tabelle mit genau dem 'pairs'-Abschnitt, dessen erkannte
    Überschrift zum Platzhalternamen passt (z.B. TABLE:ALLGEMEINE_DATEN ->
    Abschnitt "Allgemeine Daten"). Jede zweite Zeile wird wie im
    Original-Layout hellblau eingefärbt. Wird kein passender Abschnitt
    gefunden, bleibt die Tabelle leer (Platzhalterzeile wird entfernt,
    keine erfundenen Zeilen)."""
    marker_text = "{{TABLE:" + marker_key + "}}"
    marker_row, marker_idx = None, None
    for idx, row in enumerate(table.rows):
        if marker_text in "".join(c.text for c in row.cells):
            marker_row, marker_idx = row, idx
            break
    if marker_row is None:
        return

    from docx.table import _Row
    key_norm = marker_key.lower().replace("_", " ")
    matching = None
    for s in sections:
        if s["type"] == "pairs" and (
            s["heading"].lower() == key_norm or key_norm in s["heading"].lower()
            or s["heading"].lower() in key_norm
        ):
            matching = s
            break

    template_tr = marker_row._tr
    if matching:
        for row_i, (label, value) in enumerate(matching["items"]):
            new_tr = deepcopy(marker_row._tr)
            template_tr.addnext(new_tr)
            data_row = _Row(new_tr, table)
            for col_idx, text in enumerate([label, value]):
                if col_idx >= len(data_row.cells):
                    break
                cell = data_row.cells[col_idx]
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = ""
                    run = p.runs[0] if p.runs else p.add_run("")
                    run.text = text
                if row_i % 2 == 0:
                    _shade_cell(cell, shade_color)
            template_tr = new_tr

    marker_tr = table.rows[marker_idx]._tr
    marker_tr.getparent().remove(marker_tr)


def _fill_data_table(table, marker_key: str, sections: list):
    """Füllt EINE Tabelle mit allen 'pairs'-Abschnitten aus sections.
    Abschnittsüberschriften werden als fett hervorgehobene Zeile eingefügt,
    danach folgen die Label/Wert-Zeilen des Abschnitts."""
    marker_text = "{{TABLE:" + marker_key + "}}"
    marker_row, marker_idx = None, None
    for idx, row in enumerate(table.rows):
        if marker_text in "".join(c.text for c in row.cells):
            marker_row, marker_idx = row, idx
            break
    if marker_row is None:
        return

    from docx.table import _Row
    template_tr = marker_row._tr
    pair_sections = [s for s in sections if s["type"] == "pairs"]

    for section in pair_sections:
        # Überschriftszeile
        new_tr = deepcopy(template_tr)
        template_tr.addnext(new_tr)
        heading_row = _Row(new_tr, table)
        for p in heading_row.cells[0].paragraphs:
            for run in p.runs:
                run.text = ""
            run = p.runs[0] if p.runs else p.add_run("")
            run.text = section["heading"]
            run.bold = True
        for extra_cell in heading_row.cells[1:]:
            for p in extra_cell.paragraphs:
                for run in p.runs:
                    run.text = ""
        template_tr = new_tr

        for label, value in section["items"]:
            new_tr = deepcopy(marker_row._tr)
            template_tr.addnext(new_tr)
            data_row = _Row(new_tr, table)
            for col_idx, text in enumerate([label, value]):
                if col_idx >= len(data_row.cells):
                    break
                cell = data_row.cells[col_idx]
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.text = ""
                    run = p.runs[0] if p.runs else p.add_run("")
                    run.text = text
                    run.bold = False
            template_tr = new_tr

    marker_tr = table.rows[marker_idx]._tr
    marker_tr.getparent().remove(marker_tr)


def _fill_list_table(table, marker_key: str, sections: list):
    marker_text = "{{TABLE:" + marker_key + "}}"
    marker_row, marker_idx = None, None
    for idx, row in enumerate(table.rows):
        if marker_text in "".join(c.text for c in row.cells):
            marker_row, marker_idx = row, idx
            break
    if marker_row is None:
        return

    from docx.table import _Row
    template_tr = marker_row._tr
    # Nur Listen-Abschnitte übernehmen, deren erkannte Überschrift zum
    # Platzhalternamen passt (z.B. "Lieferumfang" -> TABLE:LIEFERUMFANG).
    # So werden keine fachfremden Listen (z.B. Anschriften) versehentlich
    # mit eingesammelt.
    key_norm = marker_key.lower().replace("_", "")
    matching_sections = [
        s for s in sections
        if s["type"] == "list" and key_norm[:6] in s["heading"].lower().replace(" ", "")
    ]
    items = [item for s in matching_sections for item in s["items"]]

    for item in items:
        new_tr = deepcopy(template_tr)
        template_tr.addnext(new_tr)
        row = _Row(new_tr, table)
        cell = row.cells[0]
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
            run = p.runs[0] if p.runs else p.add_run("")
            run.text = item
        template_tr = new_tr

    marker_tr = table.rows[marker_idx]._tr
    marker_tr.getparent().remove(marker_tr)


def find_matching_pairs_items(sections: list, marker_key: str):
    """Sucht den 'pairs'-Abschnitt, dessen Überschrift zum Marker passt
    (z.B. Marker 'ALLGEMEINE_DATEN' -> Abschnitt 'Allgemeine Daten') und
    gibt dessen [(label, value), ...] zurück, oder [] falls keiner passt.
    Reine Datenfunktion, unabhängig von docx - wird von build_draft()
    (Word) und build_pdf_native() (direktes PDF) gleichermaßen genutzt."""
    key_norm = marker_key.lower().replace("_", " ")
    for s in sections:
        if s["type"] == "pairs" and (
            s["heading"].lower() == key_norm or key_norm in s["heading"].lower()
            or s["heading"].lower() in key_norm
        ):
            return s["items"]
    return []


def find_matching_list_items(sections: list, marker_key: str):
    """Sucht 'list'-Abschnitte, deren Überschrift zum Marker passt (z.B.
    Marker 'LIEFERUMFANG' -> Abschnitt 'Lieferumfang') und gibt deren
    zusammengefügte Einträge zurück."""
    key_norm = marker_key.lower().replace("_", "")
    matching = [
        s for s in sections
        if s["type"] == "list" and key_norm[:6] in s["heading"].lower().replace(" ", "")
    ]
    return [item for s in matching for item in s["items"]]


def resolve_categories(extracted: dict, kategorien_module: str, not_found_text: str) -> dict:
    """Löst jede feste Kategorie (aus kategorien.py/kategorien_en.py) zu
    ihrem im Quell-PDF gefundenen Wert auf (oder not_found_text). Reine
    Datenfunktion - Grundlage für {{VAL:...}} in build_draft() UND für
    build_pdf_native()."""
    try:
        import importlib
        mod = importlib.import_module(kategorien_module)
        if hasattr(mod, "ALLGEMEINE_DATEN"):
            alle_kategorien = mod.ALLGEMEINE_DATEN + mod.SCHNITTSTELLEN + mod.TECHNISCHE_DATEN
        else:
            alle_kategorien = mod.GENERAL_DATA + mod.INTERFACES + mod.TECHNICAL_DATA
    except ImportError:
        alle_kategorien = []

    normalized_lookup = {
        _normalize_label(label): value for label, value in extracted.get("key_values", {}).items()
    }
    result = {}
    for kategorie in alle_kategorien:
        result[kategorie] = normalized_lookup.get(_normalize_label(kategorie), "") or not_found_text
    return result


def _normalize_label(s: str) -> str:
    return " ".join(s.strip().lower().split())


def build_draft(
    template_path: str,
    extracted: dict,
    output_path: str,
    kategorien_module: str = "kategorien",
    not_found_text: str = NOT_FOUND,
    lieferumfang_marker: str = "LIEFERUMFANG",
    optionen_marker: str = "OPTIONEN",
):
    """Baut aus den automatisch extrahierten Daten direkt den Entwurf -
    keine manuelle Zuordnung nötig. Alle eingesetzten Werte sind wörtliche
    Kopien aus `extracted`.

    `kategorien_module`: Name des Python-Moduls mit den festen
    Kategorienlisten - "kategorien" (Deutsch) oder "kategorien_en"
    (Englisch). Muss drei Listen ALLGEMEINE_DATEN/SCHNITTSTELLEN/
    TECHNISCHE_DATEN bzw. GENERAL_DATA/INTERFACES/TECHNICAL_DATA enthalten."""
    doc = Document(template_path)

    field_mapping = {
        "TITEL": extracted.get("titel", ""),
        "SERIE": extracted.get("serie", ""),
        "PRODUKTTYP": extracted.get("produkttyp", ""),
        "BESCHREIBUNG": extracted.get("beschreibung", ""),
    }

    # {{VAL:<Kategorie>}} - feste Kategorienzeilen. Zuordnung per
    # normalisiertem Textabgleich (getrimmt, Groß-/Kleinschreibung
    # ignoriert) gegen ALLE im Quell-PDF gefundenen Label/Wert-Paare -
    # unabhängig davon, unter welcher Überschrift sie dort standen. Wird
    # keine passende Kategorie im Quell-PDF gefunden, bleibt die Zelle
    # leer/markiert - es wird nichts erfunden.
    try:
        import importlib
        mod = importlib.import_module(kategorien_module)
        if hasattr(mod, "ALLGEMEINE_DATEN"):
            alle_kategorien = mod.ALLGEMEINE_DATEN + mod.SCHNITTSTELLEN + mod.TECHNISCHE_DATEN
        else:
            alle_kategorien = mod.GENERAL_DATA + mod.INTERFACES + mod.TECHNICAL_DATA
    except ImportError:
        alle_kategorien = []

    normalized_lookup = {
        _normalize_label(label): value for label, value in extracted.get("key_values", {}).items()
    }
    for kategorie in alle_kategorien:
        gefunden = normalized_lookup.get(_normalize_label(kategorie), "")
        field_mapping["VAL:" + kategorie] = gefunden

    for p in doc.paragraphs:
        _replace_placeholder_text(p, field_mapping, not_found_text)

    for table in doc.tables:
        _fill_named_section_table(table, optionen_marker, extracted["sections"])
        _fill_list_table(table, lieferumfang_marker, extracted["sections"])
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholder_text(p, field_mapping, not_found_text)

    doc.save(output_path)


def read_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


class PdfConversionError(Exception):
    pass


def convert_to_pdf(docx_path: str, pdf_path: str):
    """Wandelt ein .docx lokal in .pdf um - kein Internet-/Serverzugriff.
    Versucht zuerst Microsoft Word (docx2pdf, benötigt installiertes Word -
    liefert die originalgetreueste Umwandlung), fällt bei Fehler auf
    LibreOffice zurück, falls das auf dem Rechner installiert ist."""
    try:
        from docx2pdf import convert as _word_convert
        _word_convert(docx_path, pdf_path)
        if Path(pdf_path).exists():
            return
    except Exception:
        pass

    # Fallback: LibreOffice (falls installiert), headless
    import subprocess
    import shutil as _shutil
    soffice = _shutil.which("soffice") or _shutil.which("libreoffice")
    if not soffice:
        raise PdfConversionError(
            "PDF-Umwandlung fehlgeschlagen: Es wurde weder Microsoft Word "
            "noch LibreOffice auf diesem Rechner gefunden. Für die "
            "PDF-Erzeugung muss eines von beiden installiert sein."
        )
    out_dir = str(Path(pdf_path).parent)
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    produced = Path(docx_path).with_suffix(".pdf").name
    produced_path = Path(out_dir) / produced
    if not produced_path.exists():
        raise PdfConversionError(f"LibreOffice-Umwandlung fehlgeschlagen: {result.stderr[:300]}")
    if str(produced_path) != pdf_path:
        produced_path.replace(pdf_path)


def build_control_report(extracted: dict) -> str:
    """Erzeugt eine Übersicht: was wurde erkannt und wörtlich übernommen.
    Dient als Kontrolle vor dem Export."""
    lines = [f"Titel erkannt: {extracted.get('titel') or NOT_FOUND}"]
    lines.append(f"Beschreibung erkannt: {'Ja' if extracted.get('beschreibung') else 'Nein'}")
    lines.append("")
    total_pairs = 0
    for s in extracted["sections"]:
        if s["type"] == "pairs":
            lines.append(f"Abschnitt „{s['heading']}“ - {len(s['items'])} Werte:")
            for label, value in s["items"]:
                lines.append(f"   {label}: {value}")
            total_pairs += len(s["items"])
        else:
            lines.append(f"Liste „{s['heading']}“ - {len(s['items'])} Einträge:")
            for item in s["items"]:
                lines.append(f"   - {item}")
        lines.append("")
    lines.append(f"Gesamt: {total_pairs} technische Werte in {len([s for s in extracted['sections'] if s['type']=='pairs'])} Abschnitten erkannt.")
    lines.append("Alle Werte oben sind wörtliche Kopien aus dem importierten PDF (keine KI-Generierung).")
    return "\n".join(lines)
