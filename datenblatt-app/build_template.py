"""Baut vorlage.docx (Deutsch) UND vorlage_en.docx (Englisch): fester
"Bauplan" fürs Ziel-Datenblatt, in beiden Sprachen mit identischem Aufbau.

- Seite 1: Banner, Serie/Series-Titel, Beschreibung, Lieferumfang/Scope of
  delivery (dynamisch)
- Feste Kategorien-Tabellen (Allgemeine Daten/General data, Schnittstellen/
  Interfaces, Technische Daten/Technical data) - linke Spalte fest aus
  kategorien.py bzw. kategorien_en.py, rechte Spalte {{VAL:<Kategorie>}}
- Optionen: bleibt dynamisch erweiterbar
- Keine Bilder außer dem festen Firmen-Banner
- Feste Fußzeile mit den echten Firmendaten (leicht unterschiedlicher
  Wortlaut je Sprache, wie in den jeweiligen Original-Datenblättern)

Ausführen mit: python3 build_template.py
Erzeugt beide Dateien neu - eine ggf. von Hand angepasste vorlage.docx
oder vorlage_en.docx vorher sichern.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from kategorien import ALLGEMEINE_DATEN, SCHNITTSTELLEN, TECHNISCHE_DATEN
from kategorien_en import GENERAL_DATA, INTERFACES, TECHNICAL_DATA
from lib import _shade_cell


def build(
    output_path,
    banner_path,
    kopf_text,
    serie_label,
    lieferumfang_label,
    lieferumfang_marker,
    tabellen,  # Liste von (Überschrift, Kategorienliste, Platzhalter-Marker)
    optionen_label,
    optionen_marker,
    fuss1_text,
    fuss2_text,
):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    doc.add_picture(banner_path, width=Cm(17))
    doc.add_paragraph()

    kopf = doc.add_paragraph()
    kopf.add_run(kopf_text).bold = True

    titelzeile = doc.add_paragraph()
    titelzeile.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titelzeile.add_run(f"{serie_label} {{{{SERIE}}}}")
    r.bold = True
    r.underline = True
    r.font.size = Pt(13)
    doc.add_paragraph()

    doc.add_paragraph("{{BESCHREIBUNG}}")
    doc.add_paragraph()

    doc.add_paragraph().add_run(lieferumfang_label)
    t_lf = doc.add_table(rows=1, cols=1)
    t_lf.autofit = True
    t_lf.rows[0].cells[0].text = "{{TABLE:" + lieferumfang_marker + "}}"
    doc.add_paragraph()

    doc.add_paragraph().add_run("{{TITEL}}").bold = True
    doc.add_paragraph().add_run("{{PRODUKTTYP}}").bold = True
    doc.add_paragraph()

    for heading, categories, marker in tabellen:
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.underline = True
        t = doc.add_table(rows=len(categories), cols=2)
        t.autofit = True
        for i, cat in enumerate(categories):
            t.rows[i].cells[0].text = cat
            t.rows[i].cells[1].text = "{{VAL:" + cat + "}}"
            if i % 2 == 0:
                for cell in t.rows[i].cells:
                    _shade_cell(cell, "DCE9F7")
        doc.add_paragraph()

    h = doc.add_paragraph()
    r = h.add_run(optionen_label)
    r.bold = True
    r.underline = True
    t_opt = doc.add_table(rows=1, cols=2)
    t_opt.autofit = True
    t_opt.rows[0].cells[0].text = "{{TABLE:" + optionen_marker + "}}"
    doc.add_paragraph()

    doc.add_paragraph()
    fuss1 = doc.add_paragraph(fuss1_text)
    fuss1.runs[0].font.size = Pt(8)
    fuss2 = doc.add_paragraph()
    r = fuss2.add_run(fuss2_text)
    r.font.size = Pt(8)

    doc.save(output_path)
    print(f"{output_path} erzeugt ({len(tabellen)} feste Tabellen)")
    for heading, categories, _ in tabellen:
        print(f"  {heading}: {len(categories)} Zeilen")


# --- Deutsch --------------------------------------------------------------
build(
    output_path="vorlage.docx",
    banner_path="assets/banner.jpg",
    kopf_text="EPS - Datenblatt",
    serie_label="Serie",
    lieferumfang_label="Lieferumfang:",
    lieferumfang_marker="LIEFERUMFANG",
    tabellen=[
        ("Allgemeine Daten", ALLGEMEINE_DATEN, "ALLGEMEINE_DATEN"),
        ("Schnittstellen", SCHNITTSTELLEN, "SCHNITTSTELLEN"),
        ("Technische Daten", TECHNISCHE_DATEN, "TECHNISCHE_DATEN"),
    ],
    optionen_label="Optionen",
    optionen_marker="OPTIONEN",
    fuss1_text="Irrtümer und Änderungen vorbehalten/Alle Wertangaben sind typische Werte",
    fuss2_text=(
        "EPS Stromversorgung GmbH\n"
        "Electronic Power Supplies\n"
        "Alter Postweg 101 86159 Augsburg\n"
        "Tel.: +49 (0) 821 570451-0\n"
        "Fax.: +49 (0) 821 570451-25\n"
        "E-mail: info@eps-germany.de\n"
        "www.eps-germany.de"
    ),
)

# --- Englisch --------------------------------------------------------------
build(
    output_path="vorlage_en.docx",
    banner_path="assets/banner_en.jpg",
    kopf_text="EPS -  Datasheet",
    serie_label="Series",
    lieferumfang_label="Scope of delivery:",
    lieferumfang_marker="SCOPEOFDELIVERY",
    tabellen=[
        ("General data", GENERAL_DATA, "GENERAL_DATA"),
        ("Interfaces", INTERFACES, "INTERFACES"),
        ("Technical data", TECHNICAL_DATA, "TECHNICAL_DATA"),
    ],
    optionen_label="Options",
    optionen_marker="OPTIONS",
    fuss1_text="Errors and changes excepted/All values are typical values",
    fuss2_text=(
        "EPS Stromversorgung GmbH\n"
        "Electronic Power Supplies\n"
        "Alter Postweg 101, 86159 Augsburg/Germany\n"
        "Tel.: +49 (0) 821 570451-0\n"
        "E-mail: sales@eps-germany.de\n"
        "www.eps-germany.de"
    ),
)
